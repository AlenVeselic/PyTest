#! python3

# pCustomInvs.py - generates custom invitations with the use of a text file full of names

import docx

txtFile = open(r"C:\Users\alene\Downloads\automate_online-materials\guests.txt")

guests = txtFile.readlines()
txtFile.close()
invitationDoc = docx.Document()

finalParagraphNumber = 4

for guest in range(len(guests)):
    invitationDoc.add_paragraph('It would be a pleasure to have the company of')
    invitationDoc.add_paragraph(guests[guest])
    invitationDoc.add_paragraph('at 11010 Memory Lane on the Evening of')
    invitationDoc.add_paragraph('April 1st')
    paraObj1 = invitationDoc.add_paragraph("at 7 o'clock")
    paraObj1.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    #invitationDoc.paragraphs[finalParagraphNumber].runs[0]
    finalParagraphNumber += 4
    
invitationDoc.save("invs.docx")
